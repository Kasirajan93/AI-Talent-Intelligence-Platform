from docx import Document


class DOCXParser:
    """
    Extract text from DOCX resumes including:

    - Paragraphs
    - Tables
    - Header
    - Footer
    """

    @staticmethod
    def extract_text(docx_path: str):

        try:

            document = Document(docx_path)

            text = []

            # -------------------------
            # Paragraphs
            # -------------------------

            for paragraph in document.paragraphs:

                if paragraph.text.strip():

                    text.append(paragraph.text.strip())

            # -------------------------
            # Tables
            # -------------------------

            for table in document.tables:

                for row in table.rows:

                    row_text = []

                    for cell in row.cells:

                        value = cell.text.strip()

                        if value:
                            row_text.append(value)

                    if row_text:

                        text.append(" | ".join(row_text))

            # -------------------------
            # Headers
            # -------------------------

            for section in document.sections:

                header = section.header

                for paragraph in header.paragraphs:

                    if paragraph.text.strip():

                        text.append(paragraph.text.strip())

            # -------------------------
            # Footers
            # -------------------------

            for section in document.sections:

                footer = section.footer

                for paragraph in footer.paragraphs:

                    if paragraph.text.strip():

                        text.append(paragraph.text.strip())

            return "\n".join(text)

        except Exception as e:

            raise Exception(f"Error reading DOCX: {e}")